import json
from docx import Document

def normalize_title_from_education(edu_str):
    if not edu_str:
        return ""
    edu_str = edu_str.lower().replace(".", "").strip()
    if "bakal" in edu_str or edu_str.startswith("bc"):
        return "Bc."
    if "inžen" in edu_str or "ing" in edu_str:
        return "Ing."
    if "doktor" in edu_str or edu_str.startswith("phd") or "ph.d" in edu_str:
        return "Ph.D."
    if "magistr" in edu_str or "mgr" in edu_str:
        return "Mgr."
    if "mudr" in edu_str:
        return "MUDr."
    return ""

def replace_in_paragraph(paragraph, mapping):
    for key, value in mapping.items():
        placeholder = f"(({key}))"
        if placeholder in paragraph.text:
            inline = paragraph.runs
            for i in range(len(inline)):
                if placeholder in inline[i].text:
                    inline[i].text = inline[i].text.replace(placeholder, str(value))

def replace_in_table(table, mapping):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph, mapping)

# === Funkce pro Streamlit ===
def fill_document(template_name, output_filename=None):
    """
    Vyplní dokument podle šablony a dat z data.json.
    Pro použití ve Streamlit aplikaci.
    
    Args:
        template_name (str): Název šablony
        output_filename (str): Název výstupního souboru (volitelné)
    """
    try:
        # Načti data
        with open("data.json", "r", encoding="utf-8") as f:
            data = json.load(f)
        
        # Získej normalizovaný titul z nejvyssi_vzdelani
        nejvyssi_vzdelani = data.get("nejvyssi_vzdelani", "")
        titul_pred = normalize_title_from_education(nejvyssi_vzdelani)
        
        # Mapování zástupců na data
        mapping = {
            "nazev_firmy": data.get("firma_nazev", ""),
            "sidlo_firmy": data.get("firma_sidlo", ""),
            "titul_pred": titul_pred,
            "jmeno": data.get("jmeno", ""),
            "prijmeni": data.get("prijmeni", ""),
            "ico": data.get("ico", ""),
            "datum_narozeni": data.get("datum_narozeni", ""),
            "nejvyssi_vzdelani": nejvyssi_vzdelani,
            "datum_nastupu": data.get("datum_nastupu", ""),
            "povolani": data.get("povolani", ""),
            # Přidej další pole dle potřeby
        }
        
        # Načti šablonu
        doc = Document(template_name)
        
        # Nahraď ve všech odstavcích
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph, mapping)
        
        # Nahraď ve všech tabulkách
        for table in doc.tables:
            replace_in_table(table, mapping)
        
        # Ulož výstup
        if output_filename is None:
            output_filename = "vyplnena.docx"
        
        doc.save(output_filename)
        return output_filename
        
    except FileNotFoundError:
        print("❌ Soubor data.json nebyl nalezen")
        return None
    except Exception as e:
        print(f"❌ Chyba při generování dokumentu: {e}")
        return None

# Kód pro spuštění pouze pokud je soubor spuštěn přímo
if __name__ == "__main__":
    # Načti data
    try:
        with open("data.json", "r", encoding="utf-8") as f:
            data = json.load(f)
        
        # Získej normalizovaný titul z nejvyssi_vzdelani
        nejvyssi_vzdelani = data.get("nejvyssi_vzdelani", "")
        titul_pred = normalize_title_from_education(nejvyssi_vzdelani)
        
        # Mapování zástupců na data
        mapping = {
            "nazev_firmy": data.get("firma_nazev", ""),
            "sidlo_firmy": data.get("firma_sidlo", ""),
            "titul_pred": titul_pred,
            "jmeno": data.get("jmeno", ""),
            "prijmeni": data.get("prijmeni", ""),
            "ico": data.get("ico", ""),
            "datum_narozeni": data.get("datum_narozeni", ""),
            "nejvyssi_vzdelani": nejvyssi_vzdelani,
            "datum_nastupu": data.get("datum_nastupu", ""),
            "povolani": data.get("povolani", ""),
            # Přidej další pole dle potřeby
        }
        
        print("=== DEBUG: MAPOVÁNÍ KLÍČŮ ===")
        for k, v in mapping.items():
            print(f"{k}: {v}")
        print("==========================\n")
        
        # Načti šablonu
        TEMPLATE = "pop_jmeno.docx"  # uprav dle skutečného názvu šablony
        OUTPUT = "vyplnena.docx"
        
        doc = Document(TEMPLATE)
        
        # Nahraď ve všech odstavcích
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph, mapping)
        
        # Nahraď ve všech tabulkách
        for table in doc.tables:
            replace_in_table(table, mapping)
        
        # Ulož výstup
        print(f"✅ Ukládám vyplněný dokument do {OUTPUT}")
        doc.save(OUTPUT)
        
    except FileNotFoundError:
        print("❌ Soubor data.json nebyl nalezen")
    except Exception as e:
        print(f"❌ Chyba při zpracování šablony: {e}")
